import logging
import io
import re
import os
import spacy
import fitz  # PyMuPDF
from django.shortcuts import get_object_or_404
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
from PIL import Image as PILImage, ImageDraw as PILImageDraw, ImageFont as PILImageFont
import pytesseract  # For OCR
from .models import PDF
from .uploadfileSezeializer import PDFUploadSerializer

# Configurez le logging
logger = logging.getLogger(__name__)

# Charger le modèle NLP de spaCy
nlp = spacy.load('en_core_web_sm')

class FileUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        logger.debug( "Received a file upload request." )
        serializer = PDFUploadSerializer( data=request.data )
        if serializer.is_valid():
            logger.debug( "Serializer is valid." )
            file_instance = serializer.save()  # Sauvegarde l'instance du fichier
            logger.debug( f"File saved at: {file_instance.file.path}" )

            # Vérification du type de fichier pour le filtrage
            if file_instance.file.path.lower().endswith( ('.pdf', '.docx') ):
                self.process_file( file_instance.file.path )
                logger.debug( "File processed successfully." )
            else:
                logger.debug( "No processing needed for non-PDF/DOCX file." )

            return Response( serializer.data, status=status.HTTP_201_CREATED )
        else:
            logger.error( f"File upload failed: {serializer.errors}" )
            return Response( serializer.errors, status=status.HTTP_400_BAD_REQUEST )

    def process_file(self, file_path):
        if file_path.lower().endswith( '.pdf' ):
            self.process_pdf( file_path )
        elif file_path.lower().endswith( '.docx' ):
            self.process_docx( file_path )

    def process_pdf(self, file_path):
        logger.debug( f"Starting PDF processing for: {file_path}" )
        try:
            doc = fitz.open( file_path )
            new_doc = fitz.open()

            for page_num in range( len( doc ) ):
                page = doc.load_page( page_num )
                new_page = new_doc.new_page( width=page.rect.width, height=page.rect.height )

                # Copier tout le contenu de la page, incluant texte et images
                new_page.show_pdf_page( page.rect, doc, page_num )

                # Traiter les images sur la page pour détecter les scans
                self.process_images_in_page( doc, page, new_page )

            filtered_file_path = file_path.replace( ".pdf", "_filtered.pdf" )
            new_doc.save( filtered_file_path )
            new_doc.close()
            doc.close()

            os.remove( file_path )
            os.rename( filtered_file_path, file_path )
            logger.debug( f"Processed PDF saved at: {filtered_file_path}" )
        except Exception as e:
            logger.error( f"Error during PDF processing: {e}" )
            raise

    def process_images_in_page(self, doc, original_page, new_page):
        images = original_page.get_images( full=True )

        for img in images:
            xref = img[0]  # L'index de l'image dans le PDF
            base_image = doc.extract_image( xref )
            image_bytes = base_image["image"]
            img = PILImage.open( io.BytesIO( image_bytes ) )

            # Appliquer l'OCR pour extraire le texte si l'image est un scan
            extracted_text = pytesseract.image_to_string( img )
            logger.debug( f"OCR extracted text: {extracted_text}" )
            filtered_text = self.filter_sensitive_content( extracted_text )

            if filtered_text != extracted_text:
                draw = PILImageDraw.Draw( img )
                boxes = pytesseract.image_to_boxes( img )

                # Remplacer chaque mot sensible par du texte filtré tout en gardant les dimensions d'origine
                filtered_words = filtered_text.split()
                original_words = extracted_text.split()

                for i, box in enumerate( boxes.splitlines() ):
                    b = box.split( ' ' )
                    x1, y1, x2, y2 = int( b[1] ), img.height - int( b[4] ), int( b[3] ), img.height - int( b[2] )

                    if i < len( filtered_words ):
                        draw.rectangle( [x1, y1, x2, y2], fill="white" )
                        draw.text( (x1, y2), filtered_words[i], fill="black", font=PILImageFont.load_default() )

                img_byte_arr = io.BytesIO()
                img.save( img_byte_arr, format='PNG' )
                img_byte_arr = img_byte_arr.getvalue()
                new_pix = fitz.Pixmap( io.BytesIO( img_byte_arr ) )

                # Insérer l'image redressée sur la page à la même position
                new_page.insert_image( new_page.rect, pixmap=new_pix )

    def process_docx(self, file_path):
        logger.debug( f"Starting DOCX processing for: {file_path}" )
        try:
            doc = Document( file_path )
            full_text = []

            for paragraph in doc.paragraphs:
                full_text.append( paragraph.text )
            text = "\n".join( full_text )

            filtered_text = self.filter_sensitive_content( text )
            if filtered_text != text:
                # Modifier le contenu dans le document
                for paragraph in doc.paragraphs:
                    paragraph.text = self.filter_sensitive_content( paragraph.text )

                # Sauvegarder le fichier DOCX filtré
                doc.save( file_path )
                logger.debug( f"Processed DOCX saved at: {file_path}" )

        except Exception as e:
            logger.error( f"Error during DOCX processing: {e}" )
            raise

    def filter_sensitive_content(self, text):
        patterns = {
            r'\b[\w.-]+@[\w.-]+\.\w+\b': 'contenu sensible',  # Emails
            r'\b\d{3}-\d{2}-\d{4}\b': 'contenu sensible',  # SSNs
            r'\b[0-9]{2}/[0-9]{2}/[0-9]{4}\b': 'contenu sensible',  # Dates
            r'\b\d{4} \d{4} \d{4} \d{4}\b': 'contenu sensible',  # Numéros de carte de crédit
            r'\b\d+(\.\d{1,2})?\s*(€|EUR|£|\$|USD)\b': 'contenu sensible',  # Montants d'argent
        }

        for pattern, replacement in patterns.items():
            text = re.sub( pattern, replacement, text )

        doc = nlp( text )
        for ent in doc.ents:
            if ent.label_ in ['MONEY', 'GPE', 'LOC', 'ORG']:
                text = text.replace( ent.text, 'contenu sensible' )

        return text